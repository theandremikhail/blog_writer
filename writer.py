import streamlit as st
import yaml
import anthropic
import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
import fitz  # PyMuPDF for PDF parsing
import pandas as pd
import io
from pathlib import Path
import json
import re
import time
from PIL import Image
import base64
import hashlib

# -------------- PASSWORD PROTECTION ----------------
# -------------- PASSWORD PROTECTION (FIXED FOR STREAMLIT CLOUD) ----------------
def check_password():
    """Returns `True` if the user had the correct password."""
    
    def password_entered():
        """Checks whether a password entered by the user is correct."""
        # Check if password field exists in session state first
        if "password_input" in st.session_state:
            # Check password against secrets
            if 'password' in st.secrets:
                if st.session_state["password_input"] == st.secrets["password"]:
                    st.session_state["password_correct"] = True
                    del st.session_state["password_input"]  # Don't store password
                else:
                    st.session_state["password_correct"] = False
            else:
                st.error("Password not configured in secrets. Please add it to Streamlit Cloud settings.")
                st.session_state["password_correct"] = False

    # First run or password not correct
    if "password_correct" not in st.session_state:
        # Show login form
        st.markdown("""
        <div style="background: linear-gradient(135deg, #2c3e50 0%, #34495e 100%); 
                    padding: 3rem; border-radius: 10px; text-align: center; 
                    margin: 2rem auto; max-width: 500px;">
            <h1 style="color: white; margin-bottom: 0.5rem;">AIvan Login</h1>
            <p style="color: rgba(255,255,255,0.8); font-size: 1.1rem;">
                The Marketing Junction's AI Blog Writer
            </p>
        </div>
        """, unsafe_allow_html=True)
        
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            password = st.text_input(
                "Password", 
                type="password", 
                key="password_input",
                placeholder="Enter your password"
            )
            st.markdown("""
            <style>
                .stTextInput > div > div > input {
                    text-align: center;
                }
            </style>
            """, unsafe_allow_html=True)
            
            if st.button("Login", use_container_width=True):
                password_entered()
                if "password_correct" in st.session_state:
                    st.rerun()
        
        if "password_correct" in st.session_state and not st.session_state["password_correct"]:
            st.error("😕 Password incorrect. Please try again.")
        
        return False
    
    # Password correct
    elif not st.session_state["password_correct"]:
        # Show login form again
        st.markdown("""
        <div style="background: linear-gradient(135deg, #2c3e50 0%, #34495e 100%); 
                    padding: 3rem; border-radius: 10px; text-align: center; 
                    margin: 2rem auto; max-width: 500px;">
            <h1 style="color: white; margin-bottom: 0.5rem;">AIvan Login</h1>
            <p style="color: rgba(255,255,255,0.8); font-size: 1.1rem;">
                The Marketing Junction's AI Blog Writer
            </p>
        </div>
        """, unsafe_allow_html=True)
        
        col1, col2, col3 = st.columns([1, 2, 1])
        with col2:
            password = st.text_input(
                "Password", 
                type="password", 
                key="password_input",
                placeholder="Enter your password"
            )
            if st.button("Login", use_container_width=True):
                password_entered()
                if "password_correct" in st.session_state:
                    st.rerun()
        
        st.error("😕 Password incorrect. Please try again.")
        return False
    
    else:
        # Password correct
        return True

# -------------- MAIN APP STARTS HERE ----------------
# Set page config first (must be the first Streamlit command)
st.set_page_config(
    page_title="AIvan, The Marketing Junction", 
    layout="wide",
    initial_sidebar_state="expanded"
)

# Check password before showing the main app
if not check_password():
    st.stop()  # Stop execution here if password is incorrect

# -------------- CONFIG ----------------
ANTHROPIC_KEY = st.secrets["api_keys"]["anthropic_api_key"]
anthropic_client = anthropic.Anthropic(api_key=ANTHROPIC_KEY)

def load_client_config(client_name):
    with open(f"clients/{client_name}.yaml", "r") as file:
        return yaml.safe_load(file)

# -------------- HISTORY MANAGEMENT ----------------
def save_to_history(title, articles, keywords, timestamp):
    """Save generated blog to history"""
    if 'blog_history' not in st.session_state:
        st.session_state.blog_history = []
    
    history_entry = {
        'timestamp': timestamp,
        'title': title,
        'articles': articles,
        'keywords': keywords,
        'id': len(st.session_state.blog_history)
    }
    
    st.session_state.blog_history.insert(0, history_entry)  # Add to beginning
    
    # Keep only last 10 entries
    if len(st.session_state.blog_history) > 10:
        st.session_state.blog_history = st.session_state.blog_history[:10]

# -------------- FILE PROCESSING ----------------
def process_uploaded_file(uploaded_file):
    """Process uploaded file and extract text content"""
    if uploaded_file is None:
        return ""
    
    file_extension = uploaded_file.name.split('.')[-1].lower()
    
    try:
        if file_extension == 'pdf':
            # Process PDF
            pdf_document = fitz.open(stream=uploaded_file.read(), filetype="pdf")
            text = ""
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                text += page.get_text()
            pdf_document.close()
            return text
        
        elif file_extension == 'docx':
            # Process DOCX
            doc = Document(uploaded_file)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        
        elif file_extension == 'txt':
            # Process TXT
            return str(uploaded_file.read(), "utf-8")
        
        elif file_extension in ['csv', 'xlsx', 'xls']:
            # Process spreadsheet files
            if file_extension == 'csv':
                df = pd.read_csv(uploaded_file)
            else:
                df = pd.read_excel(uploaded_file)
            
            # Convert DataFrame to text summary
            text = f"Data Summary:\n"
            text += f"Shape: {df.shape[0]} rows, {df.shape[1]} columns\n"
            text += f"Columns: {', '.join(df.columns.tolist())}\n\n"
            text += "Sample Data:\n"
            text += df.head().to_string()
            
            # Add basic statistics for numeric columns
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                text += "\n\nNumeric Statistics:\n"
                text += df[numeric_cols].describe().to_string()
            
            return text
        
        else:
            st.error(f"Unsupported file type: {file_extension}")
            return ""
    
    except Exception as e:
        st.error(f"Error processing file: {str(e)}")
        return ""

# -------------- LOGO PROCESSING ----------------
def process_logo(uploaded_logo):
    """Process uploaded logo for display and storage"""
    if uploaded_logo is None:
        return None
    
    try:
        # Open and process the image
        img = Image.open(uploaded_logo)
        
        # Convert to RGB if necessary
        if img.mode != 'RGB':
            img = img.convert('RGB')
        
        # Resize if too large (maintain aspect ratio)
        max_width = 300
        if img.width > max_width:
            ratio = max_width / img.width
            new_height = int(img.height * ratio)
            img = img.resize((max_width, new_height), Image.Resampling.LANCZOS)
        
        # Save to bytes
        img_bytes = io.BytesIO()
        img.save(img_bytes, format='PNG')
        img_bytes.seek(0)
        
        return img_bytes
    except Exception as e:
        st.error(f"Error processing logo: {str(e)}")
        return None

# -------------- TITLE GENERATION ----------------
def generate_title_only(topic, client_cfg, custom_keywords=""):
    """Generate only a title for the given topic"""
    base_keywords = client_cfg.get("keywords", [])
    if custom_keywords:
        additional_keywords = [kw.strip().lower() for kw in custom_keywords.split(",") if kw.strip()]
        for kw in additional_keywords:
            if kw.lower() not in [k.lower() for k in base_keywords]:
                base_keywords.append(kw)
    keywords = ", ".join(base_keywords)
    
    # Add variation to avoid repetition
    variation = st.session_state.get('title_generation_count', 0) % 3
    style_hints = [
        "Make it compelling and action-oriented",
        "Focus on the value and benefits",
        "Make it intriguing and thought-provoking"
    ]
    
    prompt = f'''
Generate ONLY a compelling, SEO-friendly blog title for this topic: "{topic}"

Requirements:
- Professional and engaging
- Incorporate relevant keywords naturally
- Clear and specific
- 8-15 words long
- Should work for a recruitment/HR industry blog
- Keywords to consider: {keywords}
- Style: {style_hints[variation]}

Respond with ONLY the title, nothing else. No explanation, no "Title:" prefix, just the title text.
'''
    
    try:
        response = anthropic_client.messages.create(
            model="claude-3-5-sonnet-20241022",
            max_tokens=100,
            temperature=0.9,  # Higher temperature for more variation
            messages=[{"role": "user", "content": prompt}]
        )
        return response.content[0].text.strip()
    except Exception as e:
        st.error(f"Error generating title: {str(e)}")
        return None

# -------------- CLAUDE PROMPT HELPER ----------------
def generate_prompt(title, facts, quotes, ai_opt, client_cfg, custom_keywords="", document_content="", language="UK", word_range="750-1500", include_hiring_impact=False, generate_title=False):
    base_keywords = client_cfg.get("keywords", [])
    if custom_keywords:
        additional_keywords = [kw.strip().lower() for kw in custom_keywords.split(",") if kw.strip()]
        base_keywords_lower = [kw.lower() for kw in base_keywords]
        for kw in additional_keywords:
            if kw not in base_keywords_lower:
                base_keywords.append(kw)
    keywords = ", ".join(base_keywords)
    
    language_instruction = "UK English" if language == "UK" else "US English"
    spelling_note = "(use British spelling, 's' instead of 'z' in words like 'organisation')" if language == "UK" else "(use American spelling, 'z' instead of 's' in words like 'organization')"
    
    # Parse word range and OVERSHOOT to compensate
    try:
        min_words, max_words = map(int, word_range.split('-'))
    except:
        min_words, max_words = 750, 1500
    
    # OVERSHOOT the target to ensure we hit minimum
    target_words = max_words  # Just aim for maximum
    
    hiring_impact_section = ""
    if include_hiring_impact:
        hiring_impact_section = """
- **The Impact on Hiring**: 

Detailed section on how this affects recruitment, talent acquisition, hiring managers, employer branding, and recruitment strategies
"""
    
    # AI-FRIENDLY VERSION
    if ai_opt:
        prompt = f'''
Write a comprehensive {target_words}-word blog article in {language_instruction} {spelling_note} about: "{title}"

IMPORTANT: Write EXACTLY {target_words} words. This is a hard requirement.

FORMAT FOR AI-FRIENDLY/AEO OPTIMIZED CONTENT:

Structure Requirements:
- Use question headings (H2) formatted as **What is X?** or **How do I do Y?**
- Answer each question immediately with 1-2 clear sentences right after the heading
- Start major sections with **Key takeaway:** in bold
- Include one numbered step-by-step process somewhere in the article
- End with a FAQ section containing exactly 5 Q&A pairs
- Include a TL;DR summary at the very end

Content Must Include:
- 2-3 specific examples with real numbers/results (use actual industry data, not fictional)
- At least one "how-to" section with clear numbered steps
- Actionable tips readers can implement immediately
- Short paragraphs (2-3 sentences maximum)
- Use bullet points where helpful for scannability

Writing Style:
- Conversational and easy to scan
- Do not create quotes on your own
- Always put line breaks after every item on a list
- No jargon - explain complex terms simply
- Question-based headings throughout
- Direct answers immediately following questions
- Clear, practical, and actionable

Required Sections (use these as question-based headings):
1. **What is [topic]?** - Clear definition with immediate answer
2. **Why does [topic] matter?** - Key benefits with **Key takeaway:** statement
3. **How do you implement [topic]?** - Step-by-step numbered process
4. **What are the best practices for [topic]?** - Bullet points with actionable tips
5. **What challenges might you face?** - Common issues and solutions
{hiring_impact_section if include_hiring_impact and "**How does this impact hiring?**" or "" else ""}
6. **Frequently Asked Questions** - Exactly 5 Q&A pairs
7. **TL;DR Summary** - 3-4 bullet points summarizing key points

Keywords to incorporate naturally: {keywords}
{f"Include these facts: {facts}" if facts else ""}
{f"Include these quotes: {quotes}" if quotes else ""}
{f"Reference this material: {document_content[:500]}" if document_content else ""}

Remember: Use real examples and data only. Keep paragraphs short. Make it scannable.'''
    
    # STANDARD VERSION
    else:
        prompt = f'''
Write a comprehensive {target_words}-word blog article in {language_instruction} {spelling_note} about: "{title}"

IMPORTANT: Write EXACTLY {target_words} words. This is a hard requirement.

Include these sections:
- **[Opening/Lead Section - use a descriptive title, NOT "Introduction"]**:

Comprehensive overview with context and preview of main points and can be more than one paragraph. Name this section something relevant to the topic, not "Introduction"

- **[Main Section 1]**: 

Deep dive into first key aspect with examples and analysis and can be more than one paragraph

- **[Main Section 2]**: 

Exploration of second aspect with case studies and data and can be more than one paragraph

- **[Main Section 3]**: 

Discussion of challenges, opportunities, and solutions
{hiring_impact_section if include_hiring_impact else ""}
- **[Forward-Looking Section]**:

Future outlook and actionable takeaways (NOT a conclusion)

Requirements:
- DO NOT use the word "Introduction" as a heading
- Start with an engaging, topic-specific heading for the opening section
- Write detailed, expansive paragraphs (100-150 words each)
- Include specific examples, statistics, and expert insights throughout
- Use transitions and elaborate on every point
- Add single line after headings
- Format headings with ** for bold (e.g., **Understanding the Digital Transformation**)
- Incorporate these keywords naturally: {keywords}
{f"- Include these facts: {facts}" if facts else ""}
{f"- Include these quotes: {quotes}" if quotes else ""}
{f"- Reference this material: {document_content[:500]}" if document_content else ""}

Write the full {target_words}-word article now:'''
    
    return prompt, base_keywords

# -------------- ARTICLE GENERATION WITH RETRY LOGIC ----------------
def call_claude(prompt, max_tokens=8000, retry_count=3):
    """Call Claude with retry logic for 529 errors"""
    for attempt in range(retry_count):
        try:
            response = anthropic_client.messages.create(
                model="claude-3-5-sonnet-20241022",
                max_tokens=max_tokens,
                temperature=0.7,
                messages=[{"role": "user", "content": prompt}]
            )
            return response.content[0].text
        except Exception as e:
            error_message = str(e)
            if "529" in error_message or "overloaded" in error_message.lower():
                if attempt < retry_count - 1:
                    wait_time = (2 ** attempt) * 5  # 5, 10, 20 seconds
                    st.warning(f"API overloaded. Waiting {wait_time} seconds before retry...")
                    time.sleep(wait_time)
                    continue
                else:
                    st.error("API is overloaded. Please try again in a few minutes.")
                    return None
            else:
                st.error(f"Error calling Claude API: {error_message}")
                return None
    return None

# -------------- ARTICLE REVISION WITH FIXED COMPLETE OUTPUT ----------------
def revise_article(original_article, revision_request, language="UK", ai_friendly=False):
    """Revise article with color-coded output - blue for revised, black for retained"""
    language_instruction = "UK English" if language == "UK" else "US English"
    
    # Clean the article first
    clean_article = clean_article_for_display(original_article)
    current_words = len(clean_article.split())
    
    ai_format_note = ""
    if ai_friendly:
        ai_format_note = """
- Maintain AI-friendly format with question-based headings
- Keep paragraphs short (2-3 sentences max)
- Preserve the FAQ section and TL;DR summary
- Maintain conversational, scannable style
"""
    
    # More explicit prompt to ensure complete output
    prompt = f'''
I need you to revise a blog article. You MUST output the ENTIRE revised article, not just parts of it.

CURRENT ARTICLE ({current_words} words):
=========================================
{clean_article}
=========================================

REVISION REQUEST: {revision_request}

CRITICAL INSTRUCTIONS - READ CAREFULLY:

1. OUTPUT THE COMPLETE ARTICLE - Every single paragraph, every single section, from beginning to end
2. DO NOT use phrases like "[rest remains unchanged]" or "[continue with original]" or any similar shortcuts
3. DO NOT provide meta-commentary about what you changed
4. DO NOT truncate or abbreviate ANY part of the article

5. For CHANGED content: Wrap it with [REVISED] and [/REVISED] tags
6. For UNCHANGED content: Include it exactly as it was, without any tags

7. The output must be AT LEAST {current_words} words (can be up to {current_words + 100} words)
8. If revisions make it shorter, expand other sections to maintain word count
9. Maintain {language_instruction}
10. Keep all formatting with ** for bold headings
{ai_format_note}

EXAMPLE OF CORRECT OUTPUT:
**First Heading**
This paragraph stays the same from original.

[REVISED]This paragraph has been changed based on the revision request.[/REVISED]

**Second Heading**
Another unchanged paragraph here.

[REVISED]Another changed paragraph here with the requested modifications.[/REVISED]

This paragraph also unchanged.

NOW PROVIDE THE COMPLETE REVISED ARTICLE:
Every paragraph, every section, everything - with [REVISED] tags only around changed parts:'''
    
    revised_content = call_claude(prompt, max_tokens=8000, retry_count=3)
    
    if revised_content:
        # Check if the response seems truncated or incomplete
        if any(phrase in revised_content.lower() for phrase in [
            "rest remains", "continue with", "remaining sections", 
            "rest of the article", "continues unchanged", "[remaining",
            "would continue"
        ]):
            # Try again with even more forceful prompt
            prompt2 = f'''
The previous response was incomplete. I need the COMPLETE article.

Starting from this article:
{clean_article}

Apply this revision: {revision_request}

OUTPUT RULES:
- Write out EVERY SINGLE WORD of the complete article
- Use [REVISED][/REVISED] tags ONLY around changed parts
- Include EVERYTHING - no shortcuts, no summaries, no "rest continues" phrases
- Minimum {current_words} words

Output the FULL article now:'''
            
            revised_content = call_claude(prompt2, max_tokens=8000, retry_count=3)
        
        # Process the content to add HTML color tags
        processed_content = process_revision_colors(revised_content)
        return processed_content
    
    return None

def process_revision_colors(content):
    """Convert [REVISED] tags to HTML color spans for Streamlit markdown"""
    if not content:
        return content
    
    # Replace [REVISED] tags with blue color HTML
    content = content.replace('[REVISED]', '<span style="color: #0066CC;">')
    content = content.replace('[/REVISED]', '</span>')
    
    return content

def strip_html_tags(text):
    """Remove HTML tags from text for clean export"""
    if not text:
        return text
    
    # Remove all HTML tags
    clean_text = re.sub(r'<[^>]+>', '', text)
    return clean_text

# -------------- CLEAN ARTICLE FOR EXPORT ----------------
def clean_article_for_export(article):
    """Remove HTML color tags and other formatting for DOCX export"""
    if not article:
        return article
    
    # First strip HTML tags
    clean_article = strip_html_tags(article)
    
    # Then apply the existing cleaning
    clean_lines = []
    for line in clean_article.split('\n'):
        # Skip meta-commentary lines
        if any(x in line.lower() for x in [
            'word count:', 'total words:', '[total word', 
            'additional words', '---expanded', 'here\'s an additional',
            'to expand the article', 'words to expand'
        ]):
            continue
        
        # Skip separator lines
        if line.strip() in ['---', '___', '---EXPANDED CONTENT---']:
            continue
            
        clean_lines.append(line)
    
    return '\n'.join(clean_lines).strip()

# -------------- WORD COUNT ENFORCER ----------------
def ensure_word_count(article, min_words, max_words, language="UK", title="", facts="", quotes="", keywords=[], ai_friendly=False, include_hiring_impact=False):
    """Completely new approach - never shrink, only expand"""
    if not article:
        return article
    
    # Clean first
    clean_lines = []
    for line in article.split('\n'):
        if any(x in line.lower() for x in ['word count:', 'total words:', '[total word', 'additional words', '---expanded']):
            continue
        if line.strip() in ['---', '___', '---EXPANDED CONTENT---']:
            continue
        clean_lines.append(line)
    
    article = '\n'.join(clean_lines).strip()
    original_word_count = len(article.split())
    
    if original_word_count >= min_words:
        return article  # Already good!
    
    # Calculate exactly how many words we need
    words_needed = min_words - original_word_count + 50  # Small buffer
    
    st.warning(f"Article has {original_word_count} words. Need to add {words_needed} more words to reach {min_words} minimum...")
    
    # DIFFERENT APPROACH: Ask for expansion without meta-text
    expansion_prompt = f'''
I need you to write {words_needed} words of additional content that naturally expands on this article about "{title}".

Current article structure:
{article}

---

Write {words_needed} words of additional paragraphs that expand the existing topics.

CRITICAL RULES:
1. DO NOT include any labels like "Additional paragraph for..." or "Here's more content..."
2. DO NOT describe where content should go or what section it's for
3. DO NOT use phrases like "To expand on...", "Building on...", "Furthermore to the section on..."
4. Just write natural, flowing paragraphs as if they were always part of the article
5. Each paragraph should be 100-150 words of substantive content
6. Focus on concrete examples, data, analysis, and insights
7. Write in {language} English

Output ONLY the new paragraphs with no meta-commentary. Write {words_needed} words now:'''
    
    try:
        additional_content = call_claude(expansion_prompt, max_tokens=4000, retry_count=3)
        
        if additional_content:
            # Aggressive cleaning of any meta-text that slipped through
            clean_additions = []
            for line in additional_content.split('\n'):
                # Skip lines that are clearly meta-text
                line_lower = line.lower()
                if any(phrase in line_lower for phrase in [
                    'additional paragraph', 'additional content', 'here\'s', 'here is',
                    'to expand', 'this adds', 'adding to', 'for the', 'section:',
                    'i\'ll add', 'let me add', 'here are', 'to the section',
                    'building on', 'furthermore to', 'expanding on the'
                ]):
                    continue
                # Also skip if line starts with common meta-text patterns
                if line.strip().endswith(':') and len(line.strip()) < 50:
                    continue
                clean_additions.append(line)
            
            additional_content = '\n'.join(clean_additions).strip()
            
            # Combine original + additions
            expanded_article = article + "\n\n" + additional_content
            new_word_count = len(expanded_article.split())
            
            if new_word_count >= min_words:
                st.success(f"Successfully expanded article from {original_word_count} to {new_word_count} words!")
                return expanded_article
            else:
                # Still short? Add more with even stricter instructions
                still_needed = min_words - new_word_count + 50
                st.info(f"Still need {still_needed} more words. Adding more content...")
                
                more_content_prompt = f'''
Write exactly {still_needed} words about: "{title}"

Create natural paragraphs with concrete examples and analysis.

DO NOT write any introductory text, labels, or descriptions.
DO NOT say what section this is for.
Just write {still_needed} words of content:'''
                
                more_content = call_claude(more_content_prompt, max_tokens=2000, retry_count=3)
                
                if more_content:
                    # Clean again
                    clean_more = []
                    for line in more_content.split('\n'):
                        if not any(phrase in line.lower() for phrase in [
                            'additional', 'here', 'paragraph', 'section', 'adding',
                            'to expand', 'furthermore to', 'building on'
                        ]):
                            clean_more.append(line)
                    
                    more_content = '\n'.join(clean_more).strip()
                    
                    final_article = expanded_article + "\n\n" + more_content
                    final_count = len(final_article.split())
                    
                    if final_count >= min_words:
                        st.success(f"Final expansion successful: {final_count} words!")
                        return final_article
                    else:
                        st.error(f"Could not reach minimum. Final: {final_count} words (need {min_words})")
                        st.info("Use revision feature to request: 'Add 200-300 more words with additional examples and analysis'")
                        return final_article
                        
    except Exception as e:
        st.error(f"Expansion error: {str(e)}")
    
    return article

# -------------- CLEAN ARTICLE DISPLAY ----------------
def clean_article_for_display(article):
    """Remove any meta-text from article before displaying"""
    clean_lines = []
    skip_next = False
    
    for line in article.split('\n'):
        # Skip meta-commentary lines
        if any(x in line.lower() for x in [
            'word count:', 'total words:', '[total word', 
            'additional words', '---expanded', 'here\'s an additional',
            'to expand the article', 'words to expand'
        ]):
            skip_next = True
            continue
        
        # Skip separator lines
        if line.strip() in ['---', '___', '---EXPANDED CONTENT---']:
            continue
            
        # Skip line after meta-commentary
        if skip_next and line.strip() == '':
            skip_next = False
            continue
            
        skip_next = False
        clean_lines.append(line)
    
    return '\n'.join(clean_lines).strip()

# -------------- PROCESS BOLD TEXT ----------------
def process_bold_text(paragraph, p):
    """Process markdown bold text (**text**) in a paragraph for DOCX"""
    import re
    
    # Find all bold sections
    bold_pattern = r'\*\*([^*]+)\*\*'
    
    # Split the text by bold markers
    parts = re.split(bold_pattern, paragraph)
    
    # Clear the paragraph first
    p.clear()
    
    # Add parts with proper formatting
    for i, part in enumerate(parts):
        if i % 2 == 0:
            # Regular text
            if part:
                p.add_run(part)
        else:
            # Bold text
            run = p.add_run(part)
            run.bold = True
    
    return p

# -------------- CONVERT MARKDOWN TO DOCX ----------------
def markdown_to_docx(content, title):
    """Convert markdown content to DOCX format with proper bold text processing"""
    doc = Document()
    
    # Add title (without language marker)
    doc.add_heading(title, 0)
    
    # Process the content line by line
    lines = content.split('\n')
    current_paragraph = []
    
    for line in lines:
        line = line.strip()
        
        # Skip lines that are just "TITLE:" markers if present
        if line.startswith("TITLE:"):
            continue
            
        if not line:
            # Empty line - add accumulated paragraph if any
            if current_paragraph:
                p = doc.add_paragraph()
                paragraph_text = ' '.join(current_paragraph)
                process_bold_text(paragraph_text, p)
                current_paragraph = []
            continue
        
        # Check for headings with bold markers
        if line.startswith('**') and line.endswith('**') and not line[2:-2].strip().startswith('*'):
            if current_paragraph:
                p = doc.add_paragraph()
                paragraph_text = ' '.join(current_paragraph)
                process_bold_text(paragraph_text, p)
                current_paragraph = []
            
            # Extract heading text and level
            heading_text = line.strip('*').strip()
            
            # Determine heading level based on markers
            if heading_text.startswith('### '):
                heading_text = heading_text[4:]
                h = doc.add_heading(heading_text, level=3)
            elif heading_text.startswith('## '):
                heading_text = heading_text[3:]
                h = doc.add_heading(heading_text, level=2)
            elif heading_text.startswith('# '):
                heading_text = heading_text[2:]
                h = doc.add_heading(heading_text, level=1)
            else:
                # Default to level 2 for bold lines that look like headings
                h = doc.add_heading(heading_text, level=2)
            
            h.runs[0].bold = True
        # Also check for headings without bold markers (fallback)
        elif line.startswith('### '):
            if current_paragraph:
                p = doc.add_paragraph()
                paragraph_text = ' '.join(current_paragraph)
                process_bold_text(paragraph_text, p)
                current_paragraph = []
            h = doc.add_heading(line[4:], level=3)
            h.runs[0].bold = True
        elif line.startswith('## '):
            if current_paragraph:
                p = doc.add_paragraph()
                paragraph_text = ' '.join(current_paragraph)
                process_bold_text(paragraph_text, p)
                current_paragraph = []
            h = doc.add_heading(line[3:], level=2)
            h.runs[0].bold = True
        elif line.startswith('# '):
            if current_paragraph:
                p = doc.add_paragraph()
                paragraph_text = ' '.join(current_paragraph)
                process_bold_text(paragraph_text, p)
                current_paragraph = []
            h = doc.add_heading(line[2:], level=1)
            h.runs[0].bold = True
        else:
            # Regular text - accumulate
            current_paragraph.append(line)
    
    # Add any remaining paragraph
    if current_paragraph:
        p = doc.add_paragraph()
        paragraph_text = ' '.join(current_paragraph)
        process_bold_text(paragraph_text, p)
    
    return doc

# -------------- EXPORT TO DOCX ----------------
def export_docx(title, article_uk, article_us, keywords, document_analysis=""):
    timestamp = datetime.datetime.now().strftime('%Y%m%d_%H%M%S')
    safe_title = "".join(c for c in title if c.isalnum() or c in (' ', '-', '_')).rstrip()
    safe_title = safe_title.replace(' ', '_')
    
    os.makedirs("exports", exist_ok=True)
    filenames = {}
    
    # Extract actual title if it was generated
    actual_title = title
    
    # Create UK version
    if article_uk:
        # Clean article for export (removes HTML tags)
        article_uk = clean_article_for_export(article_uk)
        
        # Extract generated title if present
        if "TITLE:" in article_uk:
            title_line = article_uk.split('\n')[0]
            if title_line.startswith("TITLE:"):
                actual_title = title_line.replace("TITLE:", "").strip()
                article_uk = '\n'.join(article_uk.split('\n')[1:])  # Remove title line from content
        
        doc_uk = markdown_to_docx(article_uk, actual_title)
        
        # Add logo if available
        if 'logo_bytes' in st.session_state:
            # Add a paragraph for the logo at the beginning
            first_paragraph = doc_uk.paragraphs[0]
            logo_paragraph = first_paragraph.insert_paragraph_before()
            logo_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Save logo temporarily and add to document
            temp_logo_path = "temp_logo.png"
            with open(temp_logo_path, 'wb') as f:
                st.session_state.logo_bytes.seek(0)
                f.write(st.session_state.logo_bytes.read())
            
            logo_run = logo_paragraph.add_run()
            logo_run.add_picture(temp_logo_path, width=Pt(150))
            
            # Clean up temp file
            os.remove(temp_logo_path)
            
            # Add some spacing after logo
            doc_uk.add_paragraph("")
        
        # Add minimal metadata at the end
        doc_uk.add_paragraph("")
        doc_uk.add_paragraph("---")
        doc_uk.add_paragraph(f"Word Count: {len(article_uk.split())}")
        doc_uk.add_paragraph(f"Keywords: {', '.join(keywords)}")
        
        filename_uk = f"exports/{safe_title}_UK_{timestamp}.docx"
        doc_uk.save(filename_uk)
        filenames['UK'] = filename_uk
    
    # Create US version
    if article_us:
        # Clean article for export (removes HTML tags)
        article_us = clean_article_for_export(article_us)
        
        # Extract generated title if present
        if "TITLE:" in article_us:
            title_line = article_us.split('\n')[0]
            if title_line.startswith("TITLE:"):
                actual_title = title_line.replace("TITLE:", "").strip()
                article_us = '\n'.join(article_us.split('\n')[1:])  # Remove title line from content
        
        doc_us = markdown_to_docx(article_us, actual_title)
        
        # Add logo if available
        if 'logo_bytes' in st.session_state:
            # Add a paragraph for the logo at the beginning
            first_paragraph = doc_us.paragraphs[0]
            logo_paragraph = first_paragraph.insert_paragraph_before()
            logo_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Save logo temporarily and add to document
            temp_logo_path = "temp_logo.png"
            with open(temp_logo_path, 'wb') as f:
                st.session_state.logo_bytes.seek(0)
                f.write(st.session_state.logo_bytes.read())
            
            logo_run = logo_paragraph.add_run()
            logo_run.add_picture(temp_logo_path, width=Pt(150))
            
            # Clean up temp file
            os.remove(temp_logo_path)
            
            # Add some spacing after logo
            doc_us.add_paragraph("")
        
        # Add minimal metadata at the end
        doc_us.add_paragraph("")
        doc_us.add_paragraph("---")
        doc_us.add_paragraph(f"Word Count: {len(article_us.split())}")
        doc_us.add_paragraph(f"Keywords: {', '.join(keywords)}")
        
        filename_us = f"exports/{safe_title}_US_{timestamp}.docx"
        doc_us.save(filename_us)
        filenames['US'] = filename_us
    
    return filenames, actual_title

# Initialize session state
if 'blog_history' not in st.session_state:
    st.session_state.blog_history = []
if 'current_articles' not in st.session_state:
    st.session_state.current_articles = {}
if 'editing_mode' not in st.session_state:
    st.session_state.editing_mode = False
if 'use_generated_title' not in st.session_state:
    st.session_state.use_generated_title = False
if 'generated_title' not in st.session_state:
    st.session_state.generated_title = ""
if 'title_generation_count' not in st.session_state:
    st.session_state.title_generation_count = 0

# Professional CSS styling
st.markdown("""
<style>
    .main-header {
        background: linear-gradient(135deg, #2c3e50 0%, #34495e 100%);
        padding: 2.5rem 2rem;
        border-radius: 8px;
        margin-bottom: 2rem;
        text-align: center;
        color: white;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
    }
    
    .main-header h1 {
        margin: 0;
        font-size: 2.8rem;
        font-weight: 300;
        letter-spacing: -1px;
    }
    
    .main-header p {
        margin: 1rem 0 0 0;
        font-size: 1.1rem;
        opacity: 0.85;
        font-weight: 300;
    }
    
    .section-header {
        background: #f8f9fa;
        padding: 1.2rem 1.5rem;
        border-radius: 6px;
        border-left: 4px solid #3498db;
        margin: 1.5rem 0;
        box-shadow: 0 1px 3px rgba(0,0,0,0.05);
    }
    
    .section-header h3 {
        margin: 0;
        color: #2c3e50;
        font-weight: 500;
    }
    
    .stButton > button {
        background: linear-gradient(135deg, #3498db 0%, #2980b9 100%);
        color: white;
        border: none;
        border-radius: 6px;
        padding: 0.8rem 2.5rem;
        font-weight: 500;
        font-size: 1rem;
        transition: all 0.2s ease;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }
    
    .stButton > button:hover {
        background: linear-gradient(135deg, #2980b9 0%, #3498db 100%);
        box-shadow: 0 4px 8px rgba(0,0,0,0.15);
    }
    
    .success-message {
        background: #d4edda;
        color: #155724;
        padding: 1.5rem;
        border-radius: 6px;
        border: 1px solid #c3e6cb;
        margin: 1.5rem 0;
        box-shadow: 0 2px 4px rgba(0,0,0,0.05);
    }
    
    .success-message h4 {
        margin-top: 0;
        font-weight: 500;
    }
    
    .metric-container {
        background: white;
        padding: 1.5rem;
        border-radius: 6px;
        border: 1px solid #e1e8ed;
        box-shadow: 0 1px 3px rgba(0,0,0,0.05);
    }
    
    .professional-info {
        background: #f8f9fa;
        padding: 1rem 1.5rem;
        border-radius: 6px;
        border-left: 3px solid #17a2b8;
        margin: 1rem 0;
    }
    
    .sidebar-section {
        background: white;
        padding: 1rem;
        border-radius: 6px;
        margin-bottom: 1rem;
        border: 1px solid #e1e8ed;
    }
    
    .history-item {
        background: #f8f9fa;
        padding: 0.8rem;
        margin: 0.5rem 0;
        border-radius: 4px;
        border-left: 3px solid #6c757d;
        cursor: pointer;
        transition: all 0.2s ease;
    }
    
    .history-item:hover {
        background: #e9ecef;
        border-left-color: #3498db;
    }
    
    .edit-section {
        background: #fff3cd;
        padding: 1.5rem;
        border-radius: 6px;
        border: 1px solid #ffeaa7;
        margin: 1.5rem 0;
    }
    
    .generated-title-box {
        background: #e3f2fd;
        padding: 1rem;
        border-radius: 6px;
        border: 1px solid #90caf9;
        margin: 0.5rem 0;
    }
</style>
""", unsafe_allow_html=True)

# Header
st.markdown("""
<div class="main-header">
    <h1>Hi, I am "AI-van"!</h1>
    <p>The Marketing Junction's advanced AI-powered blog writing tool based on the human inputs and approaches of Evan.</p>
</div>
""", unsafe_allow_html=True)

# Sidebar for configuration
with st.sidebar:
    # Logout button
    if st.button("🚪 Logout"):
        for key in list(st.session_state.keys()):
            del st.session_state[key]
        st.rerun()
    
    st.markdown("---")
    
    st.markdown("### Configuration")
    
    # Client selection
    client_name = st.selectbox(
        "Client Profile",
        ["marketing_junction"],
        help="Select the client configuration profile"
    )
    
    # Logo upload (optional feature)
    st.markdown("### Company Logo (Optional)")
    uploaded_logo = st.file_uploader(
        "Upload Logo",
        type=['png', 'jpg', 'jpeg'],
        help="Upload your company logo to include in exports"
    )
    
    if uploaded_logo:
        logo_bytes = process_logo(uploaded_logo)
        if logo_bytes:
            st.image(logo_bytes, width=150)
            st.session_state.logo_bytes = logo_bytes
    
    # Language versions
    st.markdown("### Language Versions")
    generate_uk = st.checkbox("Generate UK English Version", value=False)
    generate_us = st.checkbox("Generate US English Version", value=False)
    
    # SEO Settings
    st.markdown("### SEO Settings")
    
    # Specific Keywords moved up
    extra_keywords = st.text_input(
        "Any specific Keywords",
        placeholder="keyword1, keyword2, keyword3",
        help="Comma-separated keywords to include"
    )
    
    ai_friendly = st.checkbox(
        "AI-Friendly Formatting (AEO Optimized)",
        help="Format content for AI search engines with Q&A structure"
    )
    
    # Word count setting
    st.markdown("### Content Settings")
    word_count_range = st.text_input(
        "Word Count Range",
        value="750-1500",
        help="Enter desired word count range (e.g., 750-1500)"
    )
    
    # Hiring impact section
    include_hiring_section = st.checkbox(
        "Include section on impact on hiring?",
        value=False,
        help="Add a dedicated section discussing how this topic affects recruitment and talent acquisition"
    )
    
    # Export options
    st.markdown("### Export Options")
    export_format = st.selectbox(
        "Export Format",
        ["DOCX", "PDF", "Both"],
        help="Choose export format for generated content"
    )
    
    # Blog History
    st.markdown("### Blog History")
    if st.session_state.blog_history:
        st.markdown("<small>Click to view previous blogs:</small>", unsafe_allow_html=True)
        for entry in st.session_state.blog_history[:5]:  # Show last 5
            if st.button(f"📄 {entry['title'][:30]}...", key=f"history_{entry['id']}"):
                st.session_state.current_articles = entry['articles']
                st.session_state.loaded_from_history = True
                st.rerun()
    else:
        st.info("No history yet. Generate your first blog!")

# Main content area
col1, col2 = st.columns([2, 1])

with col1:
    st.markdown('<div class="section-header"><h3>Blog Content Setup</h3></div>', unsafe_allow_html=True)
    
    # Title section with AI generation
    st.markdown("### Blog Title")
    
    # Title input and generation controls in same row
    title_col1, title_col2, title_col3 = st.columns([4, 1.5, 0.8])
    
    with title_col1:
        if st.session_state.use_generated_title and st.session_state.generated_title:
            blog_title = st.text_input(
                "Blog Title/Topic",
                value=st.session_state.generated_title,
                placeholder="Enter your compelling blog topic here...",
                help="This will be the main title of your blog post",
                key="blog_title_input"
            )
        else:
            blog_title = st.text_input(
                "Blog Title/Topic",
                placeholder="Enter your compelling blog topic here...",
                help="This will be the main title of your blog post",
                key="blog_title_input"
            )
    
    with title_col2:
        st.markdown("<div style='height: 29px'></div>", unsafe_allow_html=True)
        if st.button("🤖 AI Title", key="generate_title_btn", use_container_width=True):
            if blog_title or st.session_state.generated_title:
                topic = blog_title if blog_title else st.session_state.generated_title
                with st.spinner("Generating..."):
                    client_cfg = load_client_config(client_name)
                    suggested_title = generate_title_only(topic, client_cfg, extra_keywords)
                    if suggested_title:
                        st.session_state.generated_title = suggested_title
                        st.session_state.title_generation_count += 1
                        st.rerun()
            else:
                st.warning("Please enter a topic first")
    
    with title_col3:
        st.markdown("<div style='height: 29px'></div>", unsafe_allow_html=True)
        if st.button("🔄", key="refresh_title_btn", use_container_width=True, help="Generate another title"):
            if st.session_state.generated_title or blog_title:
                topic = blog_title if blog_title else st.session_state.generated_title
                with st.spinner("Regenerating..."):
                    client_cfg = load_client_config(client_name)
                    suggested_title = generate_title_only(topic, client_cfg, extra_keywords)
                    if suggested_title:
                        st.session_state.generated_title = suggested_title
                        st.session_state.title_generation_count += 1
                        st.rerun()
    
    # Show generated title and checkbox to use it
    if st.session_state.generated_title:
        st.markdown('<div class="generated-title-box">', unsafe_allow_html=True)
        col_a, col_b = st.columns([4, 1])
        with col_a:
            st.markdown(f"**Generated Title:** {st.session_state.generated_title}")
        with col_b:
            use_title = st.checkbox("Use this title", key="use_gen_title", value=st.session_state.use_generated_title)
            if use_title != st.session_state.use_generated_title:
                st.session_state.use_generated_title = use_title
                st.rerun()
        st.markdown('</div>', unsafe_allow_html=True)
    
    with st.form("blog_form"):
        # Content inputs
        st.markdown("### Content Inputs")
        
        content_col1, content_col2 = st.columns(2)
        
        with content_col1:
            pasted_facts = st.text_area(
                "Key Facts & Figures",
                height=120,
                placeholder="• Statistical data\n• Research findings\n• Important facts",
                help="Add relevant facts and statistics to include in your blog"
            )
        
        with content_col2:
            pasted_quotes = st.text_area(
                "Quotes & Original Thoughts",
                height=120,
                placeholder="• Expert quotes\n• Industry insights\n• Original perspectives",
                help="Add quotes and unique insights to enhance your content"
            )
        
        # File upload section
        st.markdown("### Supporting Documents")
        uploaded_file = st.file_uploader(
            "Upload Supporting Document",
            type=['pdf', 'docx', 'txt', 'csv', 'xlsx', 'xls'],
            help="Upload a document that Claude will analyze and use as supporting material"
        )
        
        if uploaded_file:
            st.success(f"File uploaded: {uploaded_file.name}")
            file_size = len(uploaded_file.getvalue()) / 1024  # KB
            st.info(f"File size: {file_size:.1f} KB")
        
        # Submit button
        submitted = st.form_submit_button("Generate Blog Articles", use_container_width=True)

with col2:
    st.markdown('<div class="section-header"><h3>Generation Settings</h3></div>', unsafe_allow_html=True)
    
    # Preview settings
    st.markdown("### Preview Settings")
    show_analysis = st.checkbox("Show Document Analysis", value=True)
    show_word_count = st.checkbox("Show Word Count", value=True)
    show_keywords = st.checkbox("Show Keywords Used", value=True)
    
    # Debug settings
    st.markdown("### Debug Settings")
    show_prompt_debug = st.checkbox("Show Prompt Debug Info", value=False, help="Shows the actual prompt being sent to Claude")
    
    # Generation metrics (placeholder)
    if 'generation_stats' not in st.session_state:
        st.session_state.generation_stats = {
            'total_blogs': 0,
            'total_words': 0,
            'files_processed': 0
        }
    
    st.markdown("### Session Statistics")
    stats_col1, stats_col2 = st.columns(2)
    
    with stats_col1:
        st.metric("Blogs Generated", st.session_state.generation_stats['total_blogs'])
        st.metric("Files Processed", st.session_state.generation_stats['files_processed'])
    
    with stats_col2:
        st.metric("Total Words", st.session_state.generation_stats['total_words'])
        st.metric("Avg Words/Blog", 
                 st.session_state.generation_stats['total_words'] // max(1, st.session_state.generation_stats['total_blogs']))

# -------------- MAIN EXECUTION ----------------
# Handle title generation and blog generation
if submitted and (blog_title or (st.session_state.use_generated_title and st.session_state.generated_title)):
    if not (generate_uk or generate_us):
        st.error("Please select at least one language version to generate.")
    else:
        # Use generated title if checkbox is checked
        if st.session_state.use_generated_title and st.session_state.generated_title:
            blog_title = st.session_state.generated_title
        
        client_cfg = load_client_config(client_name)
        
        # Process uploaded file
        document_content = ""
        if uploaded_file:
            with st.spinner("Analyzing uploaded document..."):
                document_content = process_uploaded_file(uploaded_file)
                st.session_state.generation_stats['files_processed'] += 1
        
        # Show document analysis if requested
        if show_analysis and document_content:
            st.markdown("### Document Analysis")
            with st.expander("View Document Content Summary", expanded=False):
                st.text_area("Extracted Content", document_content[:1000] + "..." if len(document_content) > 1000 else document_content, height=200)
        
        articles = {}
        
        # Debug: Show what word count is being sent
        try:
            min_words_debug, max_words_debug = map(int, word_count_range.split('-'))
        except:
            min_words_debug, max_words_debug = 750, 1500
        
        st.info(f"📊 Target: {max_words_debug} words (minimum {min_words_debug})")
        
        # Generate UK version
        if generate_uk:
            with st.spinner("Generating UK English version..."):
                full_prompt, all_keywords = generate_prompt(
                    blog_title, pasted_facts, pasted_quotes, ai_friendly, 
                    client_cfg, extra_keywords, document_content, "UK", word_count_range, 
                    include_hiring_section, generate_title=False
                )
                
                # Debug: Show part of the prompt to verify word count instructions
                if show_prompt_debug:
                    with st.expander("UK Prompt Debug Info", expanded=True):
                        st.text_area("First 1000 chars of UK prompt:", full_prompt[:1000], height=200)
                        st.info(f"Full prompt length: {len(full_prompt)} characters")
                
                article_uk = call_claude(full_prompt)
                if article_uk:
                    # Ensure word count is met - passing title now
                    article_uk = ensure_word_count(article_uk, min_words_debug, max_words_debug, "UK", 
                                                  title=blog_title, facts=pasted_facts, quotes=pasted_quotes,
                                                  keywords=all_keywords, ai_friendly=ai_friendly,
                                                  include_hiring_impact=include_hiring_section)
                    articles['UK'] = article_uk
        
        # Generate US version
        if generate_us:
            with st.spinner("Generating US English version..."):
                full_prompt, all_keywords = generate_prompt(
                    blog_title, pasted_facts, pasted_quotes, ai_friendly, 
                    client_cfg, extra_keywords, document_content, "US", word_count_range, 
                    include_hiring_section, generate_title=False
                )
                
                # Debug: Show part of the prompt to verify word count instructions
                if show_prompt_debug:
                    with st.expander("US Prompt Debug Info", expanded=True):
                        st.text_area("First 1000 chars of US prompt:", full_prompt[:1000], height=200)
                        st.info(f"Full prompt length: {len(full_prompt)} characters")
                
                article_us = call_claude(full_prompt)
                if article_us:
                    # Ensure word count is met - passing title now
                    article_us = ensure_word_count(article_us, min_words_debug, max_words_debug, "US",
                                                  title=blog_title, facts=pasted_facts, quotes=pasted_quotes,
                                                  keywords=all_keywords, ai_friendly=ai_friendly,
                                                  include_hiring_impact=include_hiring_section)
                    articles['US'] = article_us
        
        if articles:
            # Save to session state and history
            st.session_state.current_articles = articles
            st.session_state.current_keywords = all_keywords
            st.session_state.current_title = blog_title
            st.session_state.document_content = document_content
            
            # Update stats
            st.session_state.generation_stats['total_blogs'] += len(articles)
            total_words = sum(len(clean_article_for_display(article).split()) for article in articles.values())
            st.session_state.generation_stats['total_words'] += total_words
            
            # Save to history
            timestamp = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            save_to_history(blog_title, articles, all_keywords, timestamp)

elif submitted:
    st.warning("Please enter a blog title before generating articles.")

# Display generated articles or loaded from history
if st.session_state.current_articles:
    articles = st.session_state.current_articles
    
    # Success message
    st.markdown("""
    <div class="success-message">
        <h4>Blog articles ready!</h4>
        <p>Review your content below. You can request revisions before downloading.</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Revision section
    st.markdown("### Request Revisions")
    with st.container():
        revision_request = st.text_area(
            "What would you like to change?",
            placeholder="e.g., 'Make the introduction more engaging', 'Add more statistics in the second section', 'Change the tone to be more formal'",
            help="Describe specific changes you'd like to make to the article"
        )
        
        # In the revision section where buttons are handled:
        if revision_request:
            col1, col2 = st.columns(2)
            with col1:
                if 'UK' in articles and st.button("Revise UK Version", type="secondary"):
                    with st.spinner("Revising UK version..."):
                        revised_uk = revise_article(articles['UK'], revision_request, "UK", ai_friendly)
                        if revised_uk:
                            st.session_state.current_articles['UK'] = revised_uk
                            st.success("UK version revised!")
                            st.rerun()
            
            with col2:
                if 'US' in articles and st.button("Revise US Version", type="secondary"):
                    with st.spinner("Revising US version..."):
                        revised_us = revise_article(articles['US'], revision_request, "US", ai_friendly)
                        if revised_us:
                            st.session_state.current_articles['US'] = revised_us
                            st.success("US version revised!")
                            st.rerun()
    
    st.markdown("---")
    
    # Export files for download
    filenames, extracted_title = export_docx(
        st.session_state.get('current_title', 'Blog Article'),
        articles.get('UK', ''),
        articles.get('US', ''),
        st.session_state.get('current_keywords', []),
        st.session_state.get('document_content', '')
    )
    
    # Download buttons
    st.markdown("### Download Final Articles")
    download_col1, download_col2 = st.columns(2)
    
    with download_col1:
        if 'UK' in articles and 'UK' in filenames:
            with open(filenames['UK'], "rb") as file:
                st.download_button(
                    "📥 Download UK Version",
                    data=file.read(),
                    file_name=os.path.basename(filenames['UK']),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    
    with download_col2:
        if 'US' in articles and 'US' in filenames:
            with open(filenames['US'], "rb") as file:
                st.download_button(
                    "📥 Download US Version",
                    data=file.read(),
                    file_name=os.path.basename(filenames['US']),
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    
    # Preview sections
    st.markdown("---")
    st.markdown("### Article Preview")
    
    # Extract and display generated title if present
    display_title = st.session_state.get('current_title', 'Blog Article')
    if articles:
        first_article = list(articles.values())[0]
        if "TITLE:" in first_article:
            title_line = first_article.split('\n')[0]
            if title_line.startswith("TITLE:"):
                display_title = title_line.replace("TITLE:", "").strip()
                st.info(f"Generated Title: **{display_title}**")
    
    # Create tabs for different versions
    if len(articles) == 2:
        tab1, tab2 = st.tabs(["UK English", "US English"])
        
        with tab1:
            # Check if article contains color spans (indicating it's been revised)
            if '<span style="color:' in articles['UK']:
                # Display with HTML for color coding
                if show_word_count:
                    # Count words without HTML tags for accurate count
                    clean_text = re.sub(r'<[^>]+>', '', articles['UK'])
                    st.info(f"Word Count: {len(clean_text.split())} words")
                if show_keywords and 'current_keywords' in st.session_state:
                    st.info(f"Keywords: {', '.join(st.session_state.current_keywords)}")
                st.markdown(articles['UK'], unsafe_allow_html=True)
            else:
                # Clean article for display (original behavior)
                article_uk_display = clean_article_for_display(articles['UK'])
                if show_word_count:
                    st.info(f"Word Count: {len(article_uk_display.split())} words")
                if show_keywords and 'current_keywords' in st.session_state:
                    st.info(f"Keywords: {', '.join(st.session_state.current_keywords)}")
                st.markdown(article_uk_display)
        
        with tab2:
            # Check if article contains color spans (indicating it's been revised)
            if '<span style="color:' in articles['US']:
                # Display with HTML for color coding
                if show_word_count:
                    # Count words without HTML tags for accurate count
                    clean_text = re.sub(r'<[^>]+>', '', articles['US'])
                    st.info(f"Word Count: {len(clean_text.split())} words")
                if show_keywords and 'current_keywords' in st.session_state:
                    st.info(f"Keywords: {', '.join(st.session_state.current_keywords)}")
                st.markdown(articles['US'], unsafe_allow_html=True)
            else:
                # Clean article for display (original behavior)
                article_us_display = clean_article_for_display(articles['US'])
                if show_word_count:
                    st.info(f"Word Count: {len(article_us_display.split())} words")
                if show_keywords and 'current_keywords' in st.session_state:
                    st.info(f"Keywords: {', '.join(st.session_state.current_keywords)}")
                st.markdown(article_us_display)
    
    elif 'UK' in articles:
        # Check if article contains color spans (indicating it's been revised)
        if '<span style="color:' in articles['UK']:
            # Display with HTML for color coding
            if show_word_count:
                clean_text = re.sub(r'<[^>]+>', '', articles['UK'])
                st.info(f"Word Count: {len(clean_text.split())} words")
            if show_keywords and 'current_keywords' in st.session_state:
                st.info(f"Keywords: {', '.join(st.session_state.current_keywords)}")
            st.markdown(articles['UK'], unsafe_allow_html=True)
        else:
            # Clean article for display (original behavior)
            article_uk_display = clean_article_for_display(articles['UK'])
            if show_word_count:
                st.info(f"Word Count: {len(article_uk_display.split())} words")
            if show_keywords and 'current_keywords' in st.session_state:
                st.info(f"Keywords: {', '.join(st.session_state.current_keywords)}")
            st.markdown(article_uk_display)
    
    else:
        # Check if article contains color spans (indicating it's been revised)
        if '<span style="color:' in articles['US']:
            # Display with HTML for color coding
            if show_word_count:
                clean_text = re.sub(r'<[^>]+>', '', articles['US'])
                st.info(f"Word Count: {len(clean_text.split())} words")
            if show_keywords and 'current_keywords' in st.session_state:
                st.info(f"Keywords: {', '.join(st.session_state.current_keywords)}")
            st.markdown(articles['US'], unsafe_allow_html=True)
        else:
            # Clean article for display (original behavior)
            article_us_display = clean_article_for_display(articles['US'])
            if show_word_count:
                st.info(f"Word Count: {len(article_us_display.split())} words")
            if show_keywords and 'current_keywords' in st.session_state:
                st.info(f"Keywords: {', '.join(st.session_state.current_keywords)}")
            st.markdown(article_us_display)

# Footer
st.markdown("---")
st.markdown("""
<div style="text-align: center; padding: 2rem; color: #6c757d; font-size: 0.9rem;">
    <p>Powered by Claude AI | Enhanced Blog Writing Tool | The Marketing Junction</p>
</div>
""", unsafe_allow_html=True)
